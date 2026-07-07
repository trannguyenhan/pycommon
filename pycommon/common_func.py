"""
Các hàm tiện ích chung cho xử lý Word documents và chuỗi
"""
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx import Document
from docx.oxml.ns import qn
from unidecode import unidecode


def set_cell_background(cell, color):
    """Set background color for a table cell (color as hex string, e.g., 'D9D9D9')"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def set_table_borders(table):
    """Set borders for a table with single black borders"""
    tbl = table._tbl  # Truy cập XML của bảng
    tblPr = tbl.tblPr  # Truy cập thuộc tính tblPr của bảng
    tblBorders = OxmlElement('w:tblBorders')

    # Cấu hình các loại border (top, left, bottom, right, insideH, insideV)
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')      # Kiểu viền: 'single' (đơn)
        border.set(qn('w:sz'), '1')           # Độ dày (16 = 1pt)
        border.set(qn('w:space'), '0')         # Khoảng cách giữa viền
        border.set(qn('w:color'), '000000')    # Màu viền (đen)
        tblBorders.append(border)

    tblPr.append(tblBorders)  # Thêm các cài đặt border vào bảng


def append_doc(src_doc, dest_path):
    """Append content from source document to destination document"""
    dest_doc = Document(dest_path)

    for element in src_doc.element.body:
        dest_doc.element.body.append(element)

    return dest_doc


def add_heading(doc, text, style_text, FONT_NAME, FONT_SIZE, italic=False):
    """Add a heading to a document with specified formatting"""
    para = doc.add_paragraph(text, style=style_text)
    para.paragraph_format.space_after = Pt(12) # Khoảng cách sau đoạn văn
    run = para.runs[0]
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE)
    run.font.color.rgb = RGBColor(0, 0, 0)  # màu đen
    run.font.bold = True
    
    if italic:
        run.font.italic = True

    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)


def name_to_code(name: str) -> str:
    """
    Convert a name string to code format (uppercase, no diacritics, underscore separated)
    
    Args:
        name: Input name string
        
    Returns:
        Code string in format: UPPERCASE_WITH_UNDERSCORES
        
    Example:
        >>> name_to_code("Nguyễn Văn A")
        'NGUYEN_VAN_A'
    """
    if not name:
        return ''
    no_diacritics = unidecode(name) # unicode
    clean = no_diacritics.replace('-', ' ').replace('–', ' ').replace('/', ' ') # replace
    parts = clean.upper().split()
    return '_'.join(parts)

