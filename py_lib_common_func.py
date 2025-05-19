from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def set_cell_background(cell, color):
    """Set background color for a table cell (color as hex string, e.g., 'D9D9D9')"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def set_table_borders(table):
    tbl = table._tbl  # Truy cáº­p XML cá»§a báº£ng
    tblPr = tbl.tblPr  # Truy cáº­p thuá»™c tÃ­nh tblPr cá»§a báº£ng
    tblBorders = OxmlElement('w:tblBorders')

    # Cáº¥u hÃ¬nh cÃ¡c loáº¡i border (top, left, bottom, right, insideH, insideV)
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')      # Kiá»ƒu viá»n: 'single' (Ä‘Æ¡n)
        border.set(qn('w:sz'), '1')           # Äá»™ dÃ y (16 = 1pt)
        border.set(qn('w:space'), '0')         # Khoáº£ng cÃ¡ch giá»¯a viá»n
        border.set(qn('w:color'), '000000')    # MÃ u viá»n (Ä‘en)
        tblBorders.append(border)

    tblPr.append(tblBorders)  # ThÃªm cÃ¡c cÃ i Ä‘áº·t border vÃ o báº£ng

def append_doc(src_doc, dest_path):
    dest_doc = Document(dest_path)

    for element in src_doc.element.body:
        dest_doc.element.body.append(element)

    return dest_doc

def add_heading(doc, text, style_text, FONT_NAME, FONT_SIZE, italic=False):
    para = doc.add_paragraph(text, style=style_text)
    para.paragraph_format.space_after = Pt(12) # Khoáº£ng cÃ¡ch sau Ä‘oáº¡n vÄƒn
    run = para.runs[0]
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE)
    run.font.color.rgb = RGBColor(0, 0, 0)  # ðŸ‘ˆ mÃ u Ä‘en
    run.font.bold = True
    
    if italic:
        run.font.italic = True

    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

def name_to_code(name: str) -> str:
    if not name:
        return ''
    no_diacritics = unidecode(name) # unicode
    clean = no_diacritics.replace('-', ' ').replace('â€“', ' ').replace('/', ' ') # replace
    parts = clean.upper().split()
    return '_'.join(parts)
